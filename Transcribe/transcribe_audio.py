import os
import time
import assemblyai as aai
from docx import Document

# Configure AssemblyAI with your API key
aai.settings.api_key = "33d0974270ad48cea7d3e2cd072be0a8"

transcriber = aai.Transcriber()

# Function to transcribe audio file
def transcribe_audio(file_path):
    try:
        config = aai.TranscriptionConfig(speaker_labels=True)
        transcript = transcriber.transcribe(file_path, config)
        return transcript
    except Exception as e:
        print(f"Error transcribing {file_path}: {str(e)}")
        return None

# Function to format timestamps
def format_timestamp(milliseconds):
    seconds = milliseconds / 1000
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:.2f}"

# Function to save transcript to a Word document
def save_transcript_to_word(file_path, transcript):
    doc = Document()
    doc.add_heading('Transcription', level=1)

    for utterance in transcript.utterances:
        start_time = format_timestamp(utterance.start)
        utterance_text = f"Speaker {utterance.speaker} [{start_time}]: {utterance.text}"
        doc.add_paragraph(utterance_text)
    
    word_file_path = os.path.splitext(file_path)[0] + '.docx'
    doc.save(word_file_path)
    print(f"Saved transcript to {word_file_path}")

# Directory containing the m4a files
audio_files_directory = r'D:\Transcribe'  # Use raw string for the path

# Get the list of m4a files
audio_files = [f for f in os.listdir(audio_files_directory) if f.endswith('.wav')]

# Track the total duration processed
total_duration_processed = 0

# Batch process files with concurrency limit
batch_size = 5
for i in range(0, len(audio_files), batch_size):
    batch = audio_files[i:i + batch_size]
    transcriptions = []
    
    for file_name in batch:
        file_path = os.path.join(audio_files_directory, file_name)
        print(f"Processing file: {file_path}")
        transcript = transcribe_audio(file_path)
        if transcript:
            transcriptions.append((file_path, transcript))
            # Calculate the duration of the audio file from the transcript
            duration = sum([word.end - word.start for word in transcript.words]) / 1000  # in seconds
            total_duration_processed += duration / 3600  # convert to hours
            if total_duration_processed > 100:
                print("Reached the API usage limit. Please upgrade your account or wait until the next billing cycle.")
                break
    
    # Save transcriptions to Word documents
    for file_path, transcript in transcriptions:
        save_transcript_to_word(file_path, transcript)
    
    # Check if we reached the usage limit
    if total_duration_processed > 100:
        break
    
    # Respect the concurrency limit by waiting before processing the next batch
    print("Waiting for the current batch to complete...")
    time.sleep(60)  # Wait for 60 seconds to avoid hitting the concurrency limit

print("Finished processing all audio files.")
